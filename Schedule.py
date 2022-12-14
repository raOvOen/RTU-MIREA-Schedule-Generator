import openpyxl
import yaml
import datetime
import docx
from docx.shared import Pt
from docx.shared import Mm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class Pair:
    def __init__(self, name, type, teacher, place):
        self.name = name
        self.type = type
        self.teacher = teacher
        self.place = place
    def get_name(self):
        return self.name
    def get_type(self):
        return self.type
    def get_teacher(self):
        return self.teacher
    def get_place(self):
        return self.place

def set_autofit(doc):
    for t_idx, table in enumerate(doc.tables):
        doc.tables[t_idx].autofit = True
        doc.tables[t_idx].allow_autofit = True
        doc.tables[t_idx]._tblPr.xpath("./w:tblW")[0].attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
        for row_idx, r_val in enumerate(doc.tables[t_idx].rows):
            for cell_idx, c_val in enumerate(doc.tables[t_idx].rows[row_idx].cells):
                doc.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = 'auto'
                doc.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0
    return doc

def remove_row(table, del_row):
    tbl = table._tbl
    tr = del_row._tr
    tbl.remove(tr)

def trim_useless(pair_name_to_write, exclude_words):
    for i in pair_name_to_write.split():
        if i in exclude_words:
            return pair_name_to_write.split(i)[0]
        if i.split(',')[0].isdigit():
            return pair_name_to_write.split(i)[0]
    return pair_name_to_write


def parse_pair_name(pair_obj, row, week_numb, except_week_words, week_words):
    pair_name = str(pair_obj.get_name())
    pair_name_to_write = pair_name
    type_to_write = str(pair_obj.get_type())
    teacher_to_write = str(pair_obj.get_teacher())
    place_to_write = str(pair_obj.get_place())
    write_flag = True
    pair_name_list = pair_name.split()
    include_counter = 0
    i=0
    while i < len(pair_name_list):
        if pair_name_list[i] in except_week_words:
            except_weeks = []
            i+=1
            while(pair_name_list[i] not in week_words):
                for splitted_except_weeks in pair_name_list[i].split(','):
                    if splitted_except_weeks != '':
                        except_weeks.append(str(splitted_except_weeks))
                i+=1
            if (str(week_numb) in except_weeks):
                write_flag = False
                i+=1
            else:
                write_flag = True
                pair_name_to_write = ""
                for j in range(i+1, len(pair_name_list)):
                    pair_name_to_write += pair_name_list[j] + " "
                try:
                    type_to_write = (pair_obj.get_type().split("\n"))[include_counter-1]
                    teacher_to_write = (pair_obj.get_teacher().split("\n"))[include_counter-1]
                    place_to_write = (pair_obj.get_place().split("\n"))[include_counter-1]
                except Exception as e:
                    pass
                pair_name_to_write = trim_useless(pair_name_to_write, except_week_words)
                break
        elif ((pair_name_list[i].split(','))[0].isdigit()):
            include_counter += 1
            include_weeks = []
            while (pair_name_list[i] not in week_words):
                for splitted_include_weeks in pair_name_list[i].split(','):
                    if splitted_include_weeks != '':
                        include_weeks.append(str(splitted_include_weeks))
                i=i+1
            if (str(week_numb) in include_weeks):
                write_flag = True
                pair_name_to_write = ""
                for j in range(i+1, len(pair_name_list)):
                    pair_name_to_write += pair_name_list[j] + " "
                try:
                    type_to_write = (pair_obj.get_type().split("\n"))[include_counter-1]
                    teacher_to_write = (pair_obj.get_teacher().split("\n"))[include_counter-1]
                    place_to_write = (pair_obj.get_place().split("\n"))[include_counter-1]
                except Exception as e:
                    pass
                pair_name_to_write = trim_useless(pair_name_to_write, except_week_words)
                break
            else:
                write_flag = False
        i+=1

    if(write_flag):
        row.cells[3].text = pair_name_to_write
        row.cells[4].text = type_to_write
        row.cells[5].text = teacher_to_write
        row.cells[6].text = place_to_write

if __name__ == "__main__":
    print("[OvO] Simple RTU MIREA Parser. Made by raOvOen (Thx to feodoriter and V8PDL)")
    try:
        print("[*] Opening 'config.yaml' file")
        with open("config.yaml", "r", encoding="UTF-8") as ymlfile:
            cfg = yaml.safe_load(ymlfile)
    except EnvironmentError:
        print("[!] Cant find or read 'config.yaml' file")
        exit(1)
    source_path = cfg["main"]["schedule_name"]
    output_name = cfg["main"]["output_name"]
    group_name = cfg["main"]["group_name"]
    first_week_date_from_cfg = cfg["main"]["first_week_date"]
    group_name_row = cfg["additional"]["group_name_row"]
    week_day_row = cfg["additional"]["week_day_row"]
    week_day_column = cfg["additional"]["week_day_column"]
    work_week_amount = cfg["additional"]["work_week_amount"]
    work_day_amount = cfg["additional"]["work_day_amount"]
    except_week_words = cfg["additional"]["except_week_words"]
    week_words = cfg["additional"]["week_words"]
    table_style = cfg["additional"]["table_style"]
    left_margin_in_mm = int(cfg["additional"]["document_left_margin_in_mm"])
    right_margin_in_mm = int(cfg["additional"]["document_right_margin_in_mm"])
    top_margin_in_mm = int(cfg["additional"]["document_top_margin_in_mm"])
    bottom_margin_in_mm = int(cfg["additional"]["document_bottom_margin_in_mm"])
    custom_width_flag = bool(cfg["additional"]["custom_width_flag"])
    custom_width_in_mm = cfg["additional"]["custom_width"].split()
    weekdays = cfg["additional"]["weekdays"].split()
    coloring_weekdays = bool(cfg["main"]["coloring_weekdays"])
    weekdays_colors = cfg["additional"]["weekdays_colors"].split()
    except_week_words_list = except_week_words.split()
    week_words_list = week_words.split()
    source_book = openpyxl.load_workbook(source_path)
    source_worksheet = source_book.active
    print("[+] File 'config.yaml' parsed successfully")
    if "???????? ????????????" == str(source_worksheet.cell(row=week_day_row, column=week_day_column).value):
        flag = True
        pair_time = []
        temp_row = week_day_row + 2
        temp_column = week_day_column + 1
        cell_value = source_worksheet.cell(row=temp_row, column=temp_column).value
        while cell_value != 1 or flag == True:
            flag = False
            pair_period = str(source_worksheet.cell(row=temp_row, column=temp_column+1).value).replace("-",":") + " - " + str(source_worksheet.cell(row=temp_row, column=temp_column+2).value).replace("-",":")
            pair_time.append(pair_period)
            temp_row += 2
            cell_value = source_worksheet.cell(row=temp_row, column=temp_column).value
    for i in range(1, source_worksheet.max_column):
        temp_cell = source_worksheet.cell(row=group_name_row, column=i)
        if group_name in str(temp_cell.value):
            group_name_column = temp_cell.column
            break
    temp_pair_row = group_name_row + 2
    all_subjects = {}
    odd_day = [ [0]*len(pair_time) for i in range(work_day_amount)]
    even_day = [ [0]*len(pair_time) for i in range(work_day_amount)]

    for i in range(0, work_day_amount):
        for j in range(0, 2*len(pair_time)):
            temp_name = source_worksheet.cell(row=temp_pair_row, column=group_name_column).value
            temp_type = source_worksheet.cell(row=temp_pair_row, column=group_name_column+1).value
            temp_teacher = source_worksheet.cell(row=temp_pair_row, column=group_name_column+2).value
            temp_place = source_worksheet.cell(row=temp_pair_row, column=group_name_column+3).value
            if (j+1) % 2 == 1:
               odd_day[i][j-int(j/2)] = Pair(temp_name,temp_type,temp_teacher,temp_place)
            else:
               even_day[i][j - int(j/2)-1] = Pair (temp_name, temp_type, temp_teacher, temp_place)
            temp_pair_row += 1
    first_week_date=str(first_week_date_from_cfg).split()
    date = datetime.date(int(first_week_date[2]),int(first_week_date[1]),int(first_week_date[0]))
    doc = docx.Document()
    section = doc.sections[0]
    section.left_margin = Mm(left_margin_in_mm)
    section.right_margin = Mm(right_margin_in_mm)
    section.top_margin = Mm(top_margin_in_mm)
    section.bottom_margin = Mm(bottom_margin_in_mm)
    table = []
    for i in range(0,work_week_amount):
        print(f"[*] Parsing {i+1} week...")
        par = doc.add_paragraph()
        par_run = par.add_run(f"???????????? ???{i+1}\n")
        par_run.font.size = Pt(16)
        par_fmt = par.paragraph_format
        par_fmt.alignment = 1
        table.append(doc.add_table(rows=work_day_amount*len(pair_time) + 1, cols=7))
        table[i].style = table_style
        header_row = table[i].rows[0]
        header_row.cells[0].text = "???????? ??????????????"
        header_row.cells[1].text = "?????????? ????????"
        header_row.cells[2].text = "?????????? ????????????????????"
        header_row.cells[3].text = "???????????????? ????????????????????"
        header_row.cells[4].text = "?????? ????????"
        header_row.cells[5].text = "??????????????????????????"
        header_row.cells[6].text = "?????????? ????????????????????"
        for j in range(0,work_day_amount):
            for q in range(0, len(pair_time)):
                temp_row = table[i].rows[j*len(pair_time)+q + 1]
                temp_row.cells[0].text = date.strftime('%d.%m') + "\n" + weekdays[date.weekday()]
                temp_row.cells[1].text = str(q+1)
                temp_row.cells[2].text = pair_time[q]
                if((i + 1) % 2 == 1):
                    temp_day = odd_day[j][q]
                else:
                    temp_day = even_day[j][q]
                parse_pair_name(temp_day, temp_row,i+1, except_week_words_list, week_words_list)
            date = date + datetime.timedelta(days=1)
        date = date + datetime.timedelta(days=1)
        doc.add_page_break()
    print("[*] Removing empty lines...")
    for i in range(0,len(table)):
        for row in table[i].rows:
            if(row.cells[3].text == "None" or row.cells[3].text == ""):
                remove_row(table[i], row)
        if((len(table[i].rows) == 1)):
            remove_row(table[i], table[i].rows[0])
    if (custom_width_flag):
        print("[*] Setting custom width...")
        for i in range(0,len(table)):
            for row in table[i].rows:
                q = 0
                for cell in row.cells:
                    cell.width = Mm(int(custom_width_in_mm[q]))
                    q += 1
    if (not custom_width_flag):
        print("[*] Autofitting column width...")
        set_autofit(doc)
    if (coloring_weekdays):
        print("[*] Coloring the rows...")
        for i in range(0, len(table)):
            for row in table[i].rows:
                for j in range(0, len(weekdays)):
                    if (weekdays[j] in row.cells[0].text):
                        for q in range(0, 7):
                            temp_shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.replace('FFFFFF',weekdays_colors[j]).format(nsdecls('w')))
                            row.cells[q]._tc.get_or_add_tcPr().append(temp_shading_elm)
    print("[+] Parsing ended successfully!")
    doc.save(output_name)
    input("[OvO] Press Enter to exit...")



